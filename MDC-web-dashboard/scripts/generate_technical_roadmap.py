"""Generate MDC_Technical_Implementation_Roadmap.docx — a developer-facing technical roadmap."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path("/Users/sudipto/Desktop/projects/MDC-web-dashboard")
DIAGRAMS = ROOT / "scripts/diagrams/out"
OUT_DOCX = ROOT / "MDC_Technical_Implementation_Roadmap.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODEBG = "F5F5F5"


def set_cell_bg(cell, hex_):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_)
    tc_pr.append(shd)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


def code(doc, text, lang=""):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, CODEBG)
    cell.text = ""
    p = cell.paragraphs[0]
    if lang:
        run = p.add_run(f"[{lang}]\n")
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    for i, line in enumerate(text.splitlines() or [""]):
        if i > 0 or lang:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Menlo"
        r.font.size = Pt(9)
    return t


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    return t


def figure(doc, filename, caption, width_in=6.5):
    path = DIAGRAMS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


def build():
    doc = Document()

    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("MicroDataCenter Platform")
    run.font.size = Pt(22); run.bold = True; run.font.color.rgb = ACCENT
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Technical Implementation Roadmap\nPDM Observability · MeshCentral · NetFlow")
    r.font.size = Pt(14); r.italic = True; r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Audience: Backend & frontend engineers on the MDC platform\nStatus: Draft v1").font.size = Pt(10)
    doc.add_paragraph()

    # --------------------------------------------------------
    heading(doc, "1. Purpose & Scope", 1)
    para(doc,
         "This document is an implementation-oriented roadmap for three new capabilities on the MDC "
         "platform: PDM-sourced observability, MeshCentral remote connectivity, and tiered distributed "
         "NetFlow. It is written for engineers who will write the code. It is not a commercial proposal.")
    para(doc, "Each integration is described with:")
    bullet(doc, "a concrete target folder layout inside the three existing repositories,")
    bullet(doc, "the C# types, interfaces and background services to add to MicroDataCenter-WebAPI,")
    bullet(doc, "the EF Core entities and migrations required,")
    bullet(doc, "the HTTP / OData endpoint contracts exposed to the two Next.js frontends,")
    bullet(doc, "the frontend route and component additions,")
    bullet(doc, "the testing strategy needed before the slice can be considered done.")

    heading(doc, "2. Existing Codebase Baseline", 1)
    para(doc, "The three repositories the work lands in:")
    table(doc, ["Repo", "Stack", "Relevant existing surfaces"], [
        ("MDC-web-dashboard", "Next.js 14 / React 18, MSAL, Zustand, React Query, Tailwind",
         "src/app/ (routes), src/lib/mdc, src/lib/bridge, src/lib/api, src/stores, src/types"),
        ("MDC-admin-frontend", "Next.js 14 / React 18, MSAL, Zustand, React Query, Tailwind",
         "src/app/admin, src/lib/mdc, src/lib/api, src/stores, src/types"),
        ("MicroDataCenter-WebAPI", ".NET 10, C#, Controllers + OData, EF Core 10, JWT Bearer + API Key",
         "MDC.Api/Controllers, MDC.Core/Services/Providers/{PVEClient,ZeroTier,MDCDatabase,MDCEndpoint,DatacenterFactory}, MDC.Core/Models, MDC.Shared"),
    ])
    para(doc, "All three integrations add new provider modules under MDC.Core/Services/Providers/ "
              "without modifying the existing OData contracts, the WebSocket VNC relay, or the Entra ID "
              "authentication path.", italic=True)

    # --------------------------------------------------------
    heading(doc, "3. System Overview", 1)
    figure(doc, "01_system_overview.png",
           "Figure 1 — Current platform with planned additions (green boxes).", width_in=6.5)
    para(doc,
         "Three external dependencies are introduced: a self-hosted Proxmox Datacenter Manager, a "
         "self-hosted MeshCentral server, and the NetFlow ingest path. All three are reached only "
         "by the backend; browsers never talk to them directly.")

    # --------------------------------------------------------
    heading(doc, "4. Repository Impact Map", 1)
    table(doc, ["Repository", "New directories / files (abridged)"], [
        ("MicroDataCenter-WebAPI",
         "MDC.Core/Services/Providers/ProxmoxDatacenterManager/*\n"
         "MDC.Core/Services/Providers/MeshCentral/*\n"
         "MDC.Core/Services/Providers/NetFlow/*  (central)\n"
         "MDC.EdgeCollector/*  (new project for edge deployment)\n"
         "MDC.Core/Models/Observability/*\n"
         "MDC.Core/Models/Mesh/*\n"
         "MDC.Core/Models/NetFlow/*\n"
         "MDC.Core/Migrations/*  (3 new migrations)\n"
         "MDC.Api/Controllers/ObservabilityController.cs\n"
         "MDC.Api/Controllers/MeshController.cs\n"
         "MDC.Api/Controllers/FlowsController.cs\n"
         "MDC.Api/Controllers/FlowExportersController.cs"),
        ("MDC-web-dashboard",
         "src/app/workspaces/[id]/network/*  (NetFlow per-interface view)\n"
         "src/app/workspaces/[id]/vms/[vmid]/observability/*\n"
         "src/app/workspaces/[id]/vms/[vmid]/remote/*  (MeshCentral launcher)\n"
         "src/lib/api/observability.ts, flows.ts, mesh.ts\n"
         "src/components/observability/*, network/*, mesh/*\n"
         "src/hooks/useObservability.ts, useFlows.ts, useRemoteSession.ts"),
        ("MDC-admin-frontend",
         "src/app/admin/observability/*  (fleet/clusters/nodes/storage/events)\n"
         "src/app/admin/netflow/*  (exporters, explorer, physical, per-VM)\n"
         "src/app/admin/mesh/*  (device health, groups, audit)\n"
         "src/lib/api/{observability,flows,mesh}.ts\n"
         "src/components/{observability,netflow,mesh}/*"),
    ])

    # --------------------------------------------------------
    heading(doc, "5. Integration A — Proxmox Observability via PDM", 1)

    heading(doc, "5.1 Architecture", 2)
    figure(doc, "02_pdm_architecture.png",
           "Figure 2 — PDM integration. MDC backend is a thin, tier-aware, cached adapter over PDM.")

    heading(doc, "5.2 Backend Folder Layout", 2)
    code(doc,
         "MDC.Core/Services/Providers/ProxmoxDatacenterManager/\n"
         "├── PdmClient.cs                    // typed REST wrapper\n"
         "├── PdmClientOptions.cs             // base URL, token, timeouts\n"
         "├── PdmResourceResolver.cs          // MDC ↔ PDM id mapping\n"
         "├── PdmTelemetryCache.cs            // short-TTL single-flight cache\n"
         "├── PdmReconciliationService.cs     // BackgroundService\n"
         "├── PdmHealthCheck.cs\n"
         "├── Dto/\n"
         "│   ├── PdmCluster.cs / PdmNode.cs / PdmVm.cs\n"
         "│   └── PdmMetricsSample.cs\n"
         "└── Extensions/\n"
         "    └── ServiceCollectionExtensions.cs",
         lang="tree")

    heading(doc, "5.3 Key Types (skeleton)", 2)
    code(doc,
         "public interface IPdmClient\n"
         "{\n"
         "    Task<IReadOnlyList<PdmCluster>> ListClustersAsync(CancellationToken ct);\n"
         "    Task<IReadOnlyList<PdmNode>>    ListNodesAsync(string clusterId, CancellationToken ct);\n"
         "    Task<IReadOnlyList<PdmVm>>      ListVmsAsync(string clusterId, CancellationToken ct);\n"
         "    Task<PdmMetricsSample>          GetVmMetricsAsync(PdmVmRef vm, CancellationToken ct);\n"
         "    Task<PdmMetricsSample>          GetNodeMetricsAsync(PdmNodeRef node, CancellationToken ct);\n"
         "    Task<IReadOnlyList<PdmEvent>>   GetRecentEventsAsync(int take, CancellationToken ct);\n"
         "}\n\n"
         "public sealed class PdmTelemetryCache\n"
         "{\n"
         "    // TTL: metrics 5s, inventory 60s, events 10s (all configurable).\n"
         "    // Collapses concurrent identical requests into one upstream call (single-flight).\n"
         "    public Task<T> GetOrFetchAsync<T>(string key, TimeSpan ttl, Func<Task<T>> loader);\n"
         "}\n\n"
         "public sealed class PdmResourceResolver\n"
         "{\n"
         "    public Task<PdmVmRef?>   ResolveVmAsync(Guid virtualMachineEntryId);\n"
         "    public Task<Guid?>       ResolveWorkspaceAsync(PdmVmRef vm);\n"
         "    public Task ReconcileAsync(CancellationToken ct);\n"
         "}",
         lang="csharp")

    heading(doc, "5.4 EF Entity & Migration", 2)
    code(doc,
         "public sealed class PdmResourceLink            // MDC.Core/Models/Observability\n"
         "{\n"
         "    public Guid     Id               { get; set; }\n"
         "    public Guid     VirtualMachineEntryId { get; set; }\n"
         "    public string   PdmClusterId     { get; set; } = default!;\n"
         "    public string   PdmNodeId        { get; set; } = default!;\n"
         "    public int      PdmVmid          { get; set; }\n"
         "    public DateTime LastReconciledUtc{ get; set; }\n"
         "    public bool     IsStale          { get; set; }\n"
         "}\n"
         "// Unique index on (PdmClusterId, PdmVmid). FK to VirtualMachineEntry.\n"
         "// Migration: AddPdmResourceLink_20260420.cs",
         lang="csharp")

    heading(doc, "5.5 Configuration (appsettings)", 2)
    code(doc,
         "\"Pdm\": {\n"
         "  \"BaseUrl\": \"https://pdm.internal.example.com\",\n"
         "  \"ApiTokenSecretName\": \"pdm-api-token\",\n"
         "  \"MetricsCacheSeconds\": 5,\n"
         "  \"InventoryCacheSeconds\": 60,\n"
         "  \"ReconcileIntervalMinutes\": 10,\n"
         "  \"RequestTimeoutSeconds\": 10\n"
         "}",
         lang="json")

    heading(doc, "5.6 HTTP API Surface", 2)
    table(doc, ["Method & route", "Tier", "Purpose"], [
        ("GET /api/observability/vms/{vmId}/metrics", "User / Admin",
         "Live CPU, memory, disk I/O, network throughput for a single VM. User-tier: 403 if VM not in caller's workspaces."),
        ("GET /api/observability/workspaces/{wsId}/summary", "User",
         "One-row-per-VM health summary, workspace-scoped."),
        ("GET /api/observability/fleet", "Admin",
         "All clusters with aggregate health and quorum state."),
        ("GET /api/observability/clusters/{id}/nodes", "Admin", "Node list with per-node health."),
        ("GET /api/observability/nodes/{id}/metrics", "Admin", "Host-level CPU/memory/IO/network."),
        ("GET /api/observability/storage", "Admin", "Pool usage and health across the fleet."),
        ("GET /api/observability/events?take=100", "Admin", "Recent PDM task and event log."),
        ("GET /api/observability/status", "Public (auth'd)",
         "Reports PDM reachability and cache freshness — drives the degraded-mode banner."),
    ])

    heading(doc, "5.7 Frontend Changes", 2)
    para(doc, "MDC-web-dashboard", bold=True)
    bullet(doc, "New: src/app/workspaces/[id]/vms/[vmid]/observability/page.tsx — single VM live panel.")
    bullet(doc, "New: src/components/observability/VmMetricsPanel.tsx, WorkspaceHealthGrid.tsx.")
    bullet(doc, "New: src/hooks/useObservability.ts (React Query + WebSocket upgrade for live values).")
    bullet(doc, "New: src/lib/api/observability.ts (createObservabilityClient(getAccessToken)).")
    bullet(doc, "Existing VM detail page gains an Observability tab alongside Console and Network.")
    para(doc, "MDC-admin-frontend", bold=True)
    bullet(doc, "New: src/app/admin/observability/{page,fleet,clusters/[id],nodes/[id],storage,events}.tsx.")
    bullet(doc, "New: src/components/observability/FleetOverview.tsx, ClusterPanel.tsx, NodePanel.tsx, StoragePanel.tsx, EventTable.tsx, DegradedBanner.tsx.")

    heading(doc, "5.8 Testing Strategy", 2)
    bullet(doc, "Unit: PdmTelemetryCache single-flight behaviour; PdmResourceResolver mapping edge cases.")
    bullet(doc, "Integration: PdmClient against a recorded PDM API (WireMock.Net replay) — happy path, 401, 5xx, timeout.")
    bullet(doc, "Contract: ObservabilityController tier-rewrite tests — user claim without workspace → 403; user claim with workspace → filtered response.")
    bullet(doc, "Chaos: PDM unreachable for > cache TTL — verify degraded-mode banner and last-known values are returned.")
    bullet(doc, "Frontend: Playwright — VM observability tab renders, shows staleness timestamp, survives a simulated backend 503.")

    # --------------------------------------------------------
    heading(doc, "6. Integration B — MeshCentral Remote Connectivity", 1)

    heading(doc, "6.1 Topology & Session Flow", 2)
    figure(doc, "03_meshcentral.png",
           "Figure 3 — MeshCentral deployment (left) and session mint/connect sequence (right).", width_in=6.8)

    heading(doc, "6.2 Backend Folder Layout", 2)
    code(doc,
         "MDC.Core/Services/Providers/MeshCentral/\n"
         "├── MeshCentralClient.cs            // control WS wrapper (/control.ashx)\n"
         "├── MeshCentralClientOptions.cs\n"
         "├── MeshSessionBroker.cs            // authZ + login-URL mint\n"
         "├── MeshDeviceDirectory.cs          // MDC VM ↔ Mesh NodeId\n"
         "├── MeshEnrollmentHostedService.cs  // BackgroundService\n"
         "└── Dto/\n"
         "    ├── MeshDevice.cs / MeshGroup.cs\n"
         "    └── MeshSessionTicket.cs",
         lang="tree")

    heading(doc, "6.3 Key Types (skeleton)", 2)
    code(doc,
         "public interface IMeshCentralClient : IAsyncDisposable\n"
         "{\n"
         "    Task ConnectAsync(CancellationToken ct);\n"
         "    Task<IReadOnlyList<MeshDevice>> ListDevicesAsync(CancellationToken ct);\n"
         "    Task<MeshSessionTicket> MintLoginTicketAsync(string nodeId, TimeSpan ttl, CancellationToken ct);\n"
         "    Task<MeshAgentStatus>   GetAgentStatusAsync(string nodeId, CancellationToken ct);\n"
         "}\n\n"
         "public sealed class MeshSessionBroker\n"
         "{\n"
         "    // Verifies the caller has permission to VM, then asks Mesh for a login ticket\n"
         "    // and returns a signed URL. Never returns the raw control credential.\n"
         "    public Task<Uri> CreateSessionUrlAsync(ClaimsPrincipal user, Guid virtualMachineEntryId, CancellationToken ct);\n"
         "}",
         lang="csharp")

    heading(doc, "6.4 EF Entity & Migration", 2)
    code(doc,
         "public sealed class MeshDeviceLink              // MDC.Core/Models/Mesh\n"
         "{\n"
         "    public Guid     Id                    { get; set; }\n"
         "    public Guid     VirtualMachineEntryId { get; set; }\n"
         "    public string   MeshNodeId            { get; set; } = default!;\n"
         "    public string   MeshGroupId           { get; set; } = default!;\n"
         "    public DateTime FirstSeenUtc          { get; set; }\n"
         "    public DateTime? LastOnlineUtc        { get; set; }\n"
         "    public string   AgentVersion          { get; set; } = \"\";\n"
         "}\n"
         "// Unique index on MeshNodeId. FK to VirtualMachineEntry (nullable for orphaned devices).\n"
         "// Migration: AddMeshDeviceLink_20260420.cs",
         lang="csharp")

    heading(doc, "6.5 Configuration", 2)
    code(doc,
         "\"MeshCentral\": {\n"
         "  \"BaseUrl\": \"https://mesh.internal.example.com\",\n"
         "  \"ControlWsPath\": \"/control.ashx\",\n"
         "  \"ServiceUser\": \"mdc-service\",\n"
         "  \"ServiceKeySecretName\": \"meshcentral-service-key\",\n"
         "  \"LoginTicketTtlSeconds\": 60,\n"
         "  \"EnrollmentGroupPrefix\": \"mdc-ws-\"\n"
         "}",
         lang="json")

    heading(doc, "6.6 HTTP API Surface", 2)
    table(doc, ["Method & route", "Tier", "Purpose"], [
        ("POST /api/mesh/vms/{vmId}/session", "User / Admin",
         "Returns { url, expiresAt }. Backend has already verified VM access and minted a short-lived MeshCentral login URL."),
        ("GET  /api/mesh/vms/{vmId}/status", "User / Admin",
         "{ online, agentVersion, lastSeenUtc } — drives button enable/disable on the VM detail page."),
        ("POST /api/mesh/vms/{vmId}/enroll", "Admin", "Force an enrollment attempt; used for break-glass and for legacy VMs."),
        ("GET  /api/mesh/devices", "Admin", "List of all Mesh-known devices with enrolment drift vs MDC inventory."),
        ("GET  /api/mesh/sessions/audit?take=100", "Admin", "Recent session audit records surfaced from MeshCentral."),
    ])

    heading(doc, "6.7 Agent Enrollment Strategy", 2)
    bullet(doc, "New VMs: cloud-init user-data injects a one-shot installer that registers into the workspace's Mesh group. User-data template lives in MDC.Core/Services/Providers/MeshCentral/Templates/.")
    bullet(doc, "Templated VMs: optional preinstalled agent; first-boot re-registration to the correct group once MDC knows the workspace.")
    bullet(doc, "Legacy VMs: admin-triggered install via the existing VNC path — documented runbook, no automation initially.")

    heading(doc, "6.8 Frontend Changes", 2)
    para(doc, "MDC-web-dashboard", bold=True)
    bullet(doc, "VM detail page now shows two equal buttons: Remote Connect (MeshCentral) and Console (VNC). Existing VNC path untouched.")
    bullet(doc, "New: src/components/mesh/RemoteConnectButton.tsx, RemoteSessionLauncher.tsx.")
    bullet(doc, "New: src/hooks/useRemoteSession.ts — calls POST /mesh/vms/{id}/session, handles the window.open or iframe launch.")
    para(doc, "MDC-admin-frontend", bold=True)
    bullet(doc, "New: src/app/admin/mesh/{page,devices,groups,audit}.tsx.")
    bullet(doc, "New: src/components/mesh/DeviceHealthTable.tsx, EnrollmentDriftPanel.tsx, SessionAuditTable.tsx.")

    heading(doc, "6.9 Testing Strategy", 2)
    bullet(doc, "Unit: MeshSessionBroker — deny when VM not in user's workspaces, allow for admin, ticket TTL respected.")
    bullet(doc, "Integration: MeshCentralClient against a dev MeshCentral instance (docker-compose target under scripts/dev/meshcentral/).")
    bullet(doc, "End-to-end: Playwright test that clicks Remote Connect on a test VM, follows the redirect, asserts the Mesh web UI loads and the session cookie is present.")
    bullet(doc, "Negative: expired ticket replay — second redemption must 401.")

    # --------------------------------------------------------
    heading(doc, "7. Integration C — NetFlow (Tiered, Distributed)", 1)

    heading(doc, "7.1 Architecture", 2)
    figure(doc, "04_netflow_architecture.png",
           "Figure 4 — NetFlow: central collector, per-site edge collectors, central federation.", width_in=6.8)

    heading(doc, "7.2 Central Collector Folder Layout", 2)
    code(doc,
         "MDC.Core/Services/Providers/NetFlow/\n"
         "├── NetFlowListenerService.cs       // BackgroundService, UDP 2055 / 4739\n"
         "├── Parsers/\n"
         "│   ├── NetFlowV5Parser.cs\n"
         "│   ├── NetFlowV9Parser.cs          // + TemplateCache\n"
         "│   └── IpfixParser.cs\n"
         "├── FlowEnrichmentService.cs        // IP → VM / Workspace / isPhysical\n"
         "├── FlowRecordSink.cs               // batched EF Core writer\n"
         "├── FlowAggregationService.cs       // rollups\n"
         "├── FlowQueryCoordinator.cs         // federation fan-out\n"
         "├── EdgeCollectorRegistry.cs        // known edges + health\n"
         "├── Options/\n"
         "│   └── NetFlowOptions.cs\n"
         "└── Dto/\n"
         "    └── FlowRecord.cs / FlowQuery.cs / FlowPage.cs",
         lang="tree")

    heading(doc, "7.3 Edge Collector Project Layout", 2)
    code(doc,
         "MDC.EdgeCollector/                   // new .NET project, deployable as a single-file publish\n"
         "├── Program.cs                      // minimal host: listener + local store + query API\n"
         "├── LocalFlowStore.cs               // SQLite-backed ring store\n"
         "├── EdgeQueryApi.cs                 // HTTPS endpoint the central coordinator calls\n"
         "├── Authentication/\n"
         "│   └── SharedSecretHandler.cs      // rotated via MDCEndpoint channel\n"
         "└── appsettings.json                // listen port, retention, parent URL, shared secret",
         lang="tree")

    heading(doc, "7.4 Key Types (skeleton)", 2)
    code(doc,
         "public sealed record FlowRecord(\n"
         "    DateTime TsUtc,\n"
         "    string   ExporterId,\n"
         "    int      ObservationPoint,      // interface index on exporter\n"
         "    IPAddress SrcIp, int SrcPort,\n"
         "    IPAddress DstIp, int DstPort,\n"
         "    byte     Protocol,\n"
         "    long     Bytes, long Packets,\n"
         "    Guid?    VmId,                  // enriched\n"
         "    Guid?    WorkspaceId,           // enriched\n"
         "    bool     IsPhysicalInterface);\n\n"
         "public interface IFlowQueryCoordinator\n"
         "{\n"
         "    // Fans out to central store + relevant edges, merges, applies tier filter.\n"
         "    Task<FlowPage> QueryAsync(FlowQuery q, ClaimsPrincipal caller, CancellationToken ct);\n"
         "}\n\n"
         "public interface IEdgeClient\n"
         "{\n"
         "    Task<FlowPage> QueryAsync(FlowQuery q, CancellationToken ct);\n"
         "    Task<EdgeHealth> PingAsync(CancellationToken ct);\n"
         "}",
         lang="csharp")

    heading(doc, "7.5 EF Entities (central)", 2)
    code(doc,
         "public sealed class FlowExporter          // registry\n"
         "{\n"
         "    public Guid    Id            { get; set; }\n"
         "    public string  DisplayName   { get; set; } = default!;\n"
         "    public string  SourceIp      { get; set; } = default!;\n"
         "    public string  HomeCollector { get; set; } = \"central\";  // or edge id\n"
         "    public bool    IsEnabled     { get; set; }\n"
         "}\n\n"
         "public sealed class FlowRecordEntity      // only for records landing centrally\n"
         "{\n"
         "    public long    Id            { get; set; }\n"
         "    public DateTime TsUtc        { get; set; }\n"
         "    public Guid    ExporterId    { get; set; }\n"
         "    public int     ObservationPoint { get; set; }\n"
         "    public string  SrcIp         { get; set; } = default!;\n"
         "    public int     SrcPort       { get; set; }\n"
         "    public string  DstIp         { get; set; } = default!;\n"
         "    public int     DstPort       { get; set; }\n"
         "    public byte    Protocol      { get; set; }\n"
         "    public long    Bytes         { get; set; }\n"
         "    public long    Packets       { get; set; }\n"
         "    public Guid?   VmId          { get; set; }\n"
         "    public Guid?   WorkspaceId   { get; set; }\n"
         "    public bool    IsPhysicalInterface { get; set; }\n"
         "}\n"
         "// Indexes: (TsUtc DESC), (WorkspaceId, TsUtc DESC), (VmId, TsUtc DESC).",
         lang="csharp")

    heading(doc, "7.6 Tier Enforcement (central contract)", 2)
    code(doc,
         "// FlowsController.Query(...)\n"
         "var tier = _tierResolver.Resolve(User);               // Admin | WorkspaceMember\n"
         "var effective = query with {\n"
         "    WorkspaceIds = tier.Kind == TierKind.Admin\n"
         "        ? query.WorkspaceIds                            // unrestricted\n"
         "        : tier.WorkspaceIds,                            // constrained\n"
         "    IncludePhysical = tier.Kind == TierKind.Admin && query.IncludePhysical\n"
         "};\n"
         "var page = await _coordinator.QueryAsync(effective, User, ct);\n"
         "// Defensive second filter inside QueryAsync before returning,\n"
         "// in case an edge returns out-of-scope records.",
         lang="csharp")

    heading(doc, "7.7 HTTP API Surface", 2)
    table(doc, ["Method & route", "Tier", "Purpose"], [
        ("GET /api/flows", "User / Admin",
         "Paginated flow records. Server rewrites WorkspaceIds to caller's memberships for user tier."),
        ("GET /api/flows/top-talkers?scope=workspace|vm|interface&window=5m", "User / Admin", "Top-N aggregate."),
        ("GET /api/flows/by-vm/{vmId}", "User / Admin", "Convenience filter; 403 if VM not in caller workspaces."),
        ("GET /api/flows/physical", "Admin", "Physical-interface records only."),
        ("GET /api/flow-exporters", "Admin", "Registry list with home collector + health."),
        ("POST /api/flow-exporters", "Admin", "Register a new exporter."),
        ("Edge (internal) POST /internal/edge/query", "Service auth",
         "Called by FlowQueryCoordinator against each edge collector. mTLS or rotated shared secret."),
    ])

    heading(doc, "7.8 Frontend Changes", 2)
    para(doc, "MDC-web-dashboard (Workspace User)", bold=True)
    bullet(doc, "New: src/app/workspaces/[id]/network/page.tsx (interface list for this workspace).")
    bullet(doc, "New: src/app/workspaces/[id]/network/[ifaceId]/page.tsx (per-interface drilldown).")
    bullet(doc, "New: src/components/network/InterfaceList.tsx, TopTalkersCard.tsx, FlowTable.tsx, ProtocolBreakdown.tsx.")
    bullet(doc, "New: src/lib/api/flows.ts and src/hooks/useFlows.ts.")
    para(doc, "MDC-admin-frontend (Administrator)", bold=True)
    bullet(doc, "New: src/app/admin/netflow/{page,exporters,explorer,physical,vms/[id]}.tsx.")
    bullet(doc, "New: src/components/netflow/ExporterTable.tsx, FlowExplorer.tsx, PhysicalInterfacePanel.tsx, PerVmDrilldown.tsx.")

    heading(doc, "7.9 Testing Strategy", 2)
    bullet(doc, "Unit: each parser against a corpus of recorded datagrams (fixtures under tests/fixtures/netflow/).")
    bullet(doc, "Unit: FlowEnrichmentService — IP → VM mapping, IP in two workspaces, unmatched IP returns null ids.")
    bullet(doc, "Contract: tier rewriter — user tier cannot pass WorkspaceIds they do not belong to; IncludePhysical silently forced false.")
    bullet(doc, "Integration: central + one edge in docker-compose; assert federated query returns merged results with correct tier filtering.")
    bullet(doc, "Load: 1 k records/sec for 10 minutes into central store; verify query latency bounds.")

    # --------------------------------------------------------
    heading(doc, "8. Cross-Cutting Concerns", 1)

    heading(doc, "8.1 Auth & Claim Model", 2)
    bullet(doc, "Admin tier: presence of the existing Administrator role claim (confirm exact claim name with Identity team before coding).")
    bullet(doc, "Workspace-member tier: workspaces claim (array of workspace IDs) derived from existing IAM. If absent, a backend call to IWorkspaceService resolves memberships and caches for the request lifetime.")
    bullet(doc, "A single TierResolver lives at MDC.Core/Services/Security/TierResolver.cs and is consumed by ObservabilityController, FlowsController, and MeshSessionBroker — one place, three call sites.")

    heading(doc, "8.2 Configuration & Secrets", 2)
    bullet(doc, "All three integrations add top-level sections to appsettings.json (Pdm, MeshCentral, NetFlow).")
    bullet(doc, "Every secret (PDM token, MeshCentral service key, edge shared secret) is resolved through the existing secret-store abstraction, never committed.")
    bullet(doc, "Edge collectors receive their shared secret through the existing MDCEndpoint registration channel; rotation is a background job.")

    heading(doc, "8.3 Logging & Audit", 2)
    bullet(doc, "Every session mint (MeshCentral) and every admin-tier query (NetFlow physical, ObservabilityController admin endpoints) writes an audit record via the existing IAuditLog sink.")
    bullet(doc, "PDM request/response sizes are logged at Debug; failures at Warning; circuit-breaker state transitions at Info.")

    heading(doc, "8.4 Error Handling & Degraded Modes", 2)
    bullet(doc, "PDM unreachable → ObservabilityController returns 200 with X-MDC-Degraded: pdm-unreachable header and last-known cached values; UI renders the degraded banner.")
    bullet(doc, "MeshCentral unreachable → POST /session returns 503 with a body pointing the UI to the VNC fallback.")
    bullet(doc, "Edge collector unreachable → FlowQueryCoordinator returns partial results with a collectorsMissing field enumerating which edges timed out; UI shows a warning badge.")

    # --------------------------------------------------------
    heading(doc, "9. Delivery Roadmap", 1)
    figure(doc, "05_phase_dependencies.png",
           "Figure 5 — Phase dependency graph. Arrows indicate prerequisites.", width_in=6.8)
    table(doc, ["Phase", "Work", "Depends on", "Exit criterion"], [
        ("P1", "PDM stand-up, PdmClient, resource resolver, observability thin slice for one cluster", "—",
         "`GET /api/observability/fleet` returns a real cluster from PDM"),
        ("P2", "User observability panel in MDC-web-dashboard", "P1",
         "Workspace user sees live metrics for their own VMs only"),
        ("P3", "Admin observability section in MDC-admin-frontend", "P1",
         "Admin can operate Proxmox estate from MDC without visiting PDM UI"),
        ("P4", "MeshCentral server + trust broker thin slice", "—",
         "One enrolled VM reachable via Remote Connect button"),
        ("P5", "User Remote Connect UI + cloud-init enrollment", "P4",
         "Newly created VM auto-enrolls and is reachable from browser"),
        ("P6", "NetFlow central collector + tier-aware API", "P1",
         "One exporter's flows visible via /api/flows, correctly tagged"),
        ("P7", "User NetFlow UI", "P6",
         "Workspace user sees per-vNIC flow analysis, no cross-workspace leakage"),
        ("P8", "Admin NetFlow UI", "P6",
         "Admin can answer 'what is VM X talking to' and see physical uplinks"),
        ("P9", "Edge collector + federation", "P8",
         "At least one edge in production, queryable transparently"),
        ("P10", "Admin MeshCentral UI", "P5",
         "Admin lifecycle operations for Mesh inside MDC admin"),
        ("P11", "Hardening and handover", "P7, P9, P10",
         "Runbooks + training session delivered"),
    ])

    # --------------------------------------------------------
    heading(doc, "10. Local Development Setup", 1)
    bullet(doc, "docker-compose profile `dev` brings up: PDM (where available) or a mock PDM container, a MeshCentral server, a NetFlow exporter simulator (e.g. nfgen), and a single edge collector. Files under scripts/dev/.")
    bullet(doc, "`dotnet run --project MDC.Api` picks up appsettings.Development.json with dev URLs and test secrets from user-secrets.")
    bullet(doc, "Frontends: pnpm dev as today. New env vars: NEXT_PUBLIC_OBSERVABILITY_ENABLED, NEXT_PUBLIC_MESH_ENABLED, NEXT_PUBLIC_NETFLOW_ENABLED (feature gates for progressive rollout).")
    bullet(doc, "Smoke script: scripts/dev/smoke.sh — hits each new endpoint with a dev Bearer token and asserts shape.")

    heading(doc, "11. Definition of Done (per integration)", 1)
    table(doc, ["Checkpoint", "PDM Observability", "MeshCentral", "NetFlow"], [
        ("Backend provider + config", "✓", "✓", "✓"),
        ("EF migration committed and applied in CI", "✓", "✓", "✓ (central)"),
        ("HTTP contracts covered by unit + contract tests", "✓", "✓", "✓"),
        ("Tier enforcement tests (user-scoped + admin-scoped)", "✓", "✓ (authZ on session mint)", "✓"),
        ("Frontend routes behind feature flag", "✓", "✓", "✓"),
        ("End-to-end Playwright smoke", "✓ (VM panel loads)",
         "✓ (Remote Connect redirect lands on Mesh UI)",
         "✓ (user flow list filtered to own workspace)"),
        ("Degraded-mode path documented and tested", "✓ (PDM down)",
         "✓ (Mesh down → VNC fallback)",
         "✓ (edge down → partial-result banner)"),
        ("Operator runbook updated", "✓", "✓", "✓"),
    ])

    doc.save(str(OUT_DOCX))
    print(f"Saved: {OUT_DOCX}")


if __name__ == "__main__":
    build()
