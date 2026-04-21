"""Generate a client-facing proposal document for NetFlow + MeshCentral integrations."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_kv_row(table, key, value):
    row = table.add_row().cells
    row[0].text = key
    row[1].text = value
    for p in row[0].paragraphs:
        for r in p.runs:
            r.bold = True
    return row


def add_phase_table(doc, phases):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Phase", "Scope", "Duration (est.)", "Primary Deliverable"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for phase in phases:
        row = table.add_row().cells
        for i, val in enumerate(phase):
            row[i].text = val


def build():
    doc = Document()

    # ---- Default style ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MicroDataCenter Platform")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("NetFlow Monitoring, Proxmox Observability, and MeshCentral Remote Connectivity — Integration Proposal")
    sub_run.font.size = Pt(14)
    sub_run.italic = True
    sub_run.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared for: Client Review\nDocument Status: Draft for Approval\nDate: April 2026").font.size = Pt(10)

    doc.add_paragraph()

    # ---- 1. Executive Summary ----
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(
        doc,
        "This document proposes three new capabilities for the MicroDataCenter (MDC) platform: "
        "(A) tiered NetFlow-based network traffic visibility, available per-interface to workspace "
        "users and globally (including physical interfaces) to administrators; "
        "(B) MeshCentral-based browser remote connectivity to virtual machines, self-hosted and "
        "coexisting with the existing VNC console; and "
        "(C) Proxmox observability via the Proxmox Datacenter Manager (PDM), consumed by the MDC "
        "portal as a centralised telemetry source across all Proxmox clusters.",
    )
    add_para(
        doc,
        "Together, these additions give users a complete picture of their own workloads — how they "
        "behave on the network, how they consume resources, and how to reach them — while giving "
        "administrators a single operational plane across every cluster, site, and workspace.",
    )
    add_para(
        doc,
        "The plan is structured as a set of thin, independently-deployable slices so that value is "
        "delivered incrementally and any slice can be reviewed, tested, and signed off before the "
        "next begins.",
    )

    # ---- 2. Current Platform Context ----
    add_heading(doc, "2. Current Platform Context", level=1)
    add_para(doc, "The existing platform consists of three cooperating components:", bold=False)
    add_bullet(doc, "MDC-web-dashboard — end-user Next.js 14 / React 18 application; authenticates against Microsoft Entra ID (MSAL) and consumes the backend API.")
    add_bullet(doc, "MDC-admin-frontend — administrator Next.js 14 / React 18 application sharing the same tech stack and auth model as the user dashboard.")
    add_bullet(doc, "MicroDataCenter-WebAPI — .NET 10 / C# backend, OData-based controllers, Entity Framework Core 10, JWT Bearer + API Key dual auth; orchestrates Proxmox (PVE), ZeroTier overlay networking, endpoint registration, and a WebSocket-based VNC relay for VM console access.")
    add_para(
        doc,
        "Neither NetFlow collection, MeshCentral integration, nor Proxmox Datacenter Manager (PDM) "
        "integration currently exists in the codebase. All three additions fit as new provider "
        "modules inside the existing backend architecture without requiring breaking changes to "
        "OData contracts, authentication, or the existing VNC console path. PDM specifically is "
        "an operational system outside the MDC codebase — the MDC backend becomes a consumer of "
        "its API, not an author of new telemetry pipelines.",
    )

    # ---- 3. Goals & Non-Goals ----
    add_heading(doc, "3. Goals & Non-Goals", level=1)
    add_heading(doc, "3.1 Goals", level=2)
    add_bullet(doc, "Give workspace users per-interface traffic visibility for the virtual networks inside their own workspace, and nothing outside it.")
    add_bullet(doc, "Give administrators an unrestricted view across all workspaces plus the physical network interfaces of the underlying system.")
    add_bullet(doc, "Keep central storage footprint minimal by supporting distributed NetFlow collection on adjacent site devices, each with local storage.")
    add_bullet(doc, "Allow end-users to remotely connect to their VMs from any modern browser — including RDP, SSH, SFTP, and console — without installing a client.")
    add_bullet(doc, "Run MeshCentral self-hosted alongside, not instead of, the existing Proxmox VNC WebSocket relay. Both remain available per VM.")
    add_bullet(doc, "Provide Proxmox telemetry (host, cluster, VM, container, storage) inside the MDC portal, sourced from a single Proxmox Datacenter Manager instance so the MDC backend does not have to build its own multi-cluster aggregation.")
    add_bullet(doc, "Tier observability the same way NetFlow is tiered: workspace users see only their own VMs' health; administrators see fleet-wide metrics and host-level health.")
    add_bullet(doc, "Preserve the existing Entra ID single sign-on experience across all new features.")
    add_bullet(doc, "Deliver in small, reviewable slices with clear sign-off points.")

    add_heading(doc, "3.2 Non-Goals (for this phase)", level=2)
    add_bullet(doc, "Replacing the existing Proxmox VNC relay. VNC remains a first-class, supported path; MeshCentral is additive.")
    add_bullet(doc, "Packet-level (full PCAP) capture. NetFlow provides metadata-level flow records, not payload capture.")
    add_bullet(doc, "Automated threat response or IDS/IPS functionality. NetFlow output may feed such systems later but is not one itself.")
    add_bullet(doc, "Cross-workspace visibility for non-administrator users. Workspace boundaries are strictly enforced at the API layer.")
    add_bullet(doc, "Mobile applications. All new UI lives in the existing Next.js web applications.")

    # ---- 4. Part A: NetFlow ----
    add_heading(doc, "4. Part A — NetFlow Integration (Tiered, User + Administrator)", level=1)

    add_heading(doc, "4.1 What NetFlow Provides", level=2)
    add_para(
        doc,
        "NetFlow (and its successors v9 and IPFIX) is an industry-standard protocol originally "
        "developed by Cisco. Network devices — routers, switches, firewalls, and virtual switches — "
        "summarise every conversation passing through them into a \"flow record\" describing source "
        "and destination IP addresses, ports, protocol, byte/packet counts, and timestamps. These "
        "records are exported as UDP datagrams to a collector. NetFlow answers questions such as: "
        "which VMs are generating the most traffic, which external destinations are being contacted, "
        "when peaks occur, and whether any unexpected east-west communication is taking place.",
    )

    add_heading(doc, "4.2 Business Value", level=2)
    add_bullet(doc, "Workspace self-service — users diagnose their own application traffic per virtual NIC without needing to raise a support ticket.")
    add_bullet(doc, "Capacity planning — evidence-based decisions on uplink sizing and VM placement.")
    add_bullet(doc, "Troubleshooting — quickly identify which VM, virtual interface, or service is consuming bandwidth during incidents.")
    add_bullet(doc, "Security hygiene — detect lateral movement, unexpected egress, or misconfigured services early.")
    add_bullet(doc, "Chargeback / showback — attribute bandwidth usage to workspaces or tenants.")

    add_heading(doc, "4.3 Access Tiers", level=2)
    add_para(
        doc,
        "NetFlow visibility is provided in two distinct tiers. The backend enforces tier membership; "
        "the frontends render the tier the caller is entitled to.",
    )
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Audience", "Scope of Visibility", "Where Surfaced"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tiers = [
        ("Workspace User",
         "Flows observed on virtual network interfaces inside the workspaces the user is a member of. Per-interface breakdown (per vNIC / per virtual network). No visibility into other workspaces, no visibility into physical NICs.",
         "MDC-web-dashboard — new Network tab within each workspace."),
        ("Administrator",
         "All workspace flows plus flows on the physical network interfaces of the underlying MDC hosts (uplinks, management NICs, trunk ports). Cross-workspace top-talkers, global traffic maps, exporter health.",
         "MDC-admin-frontend — new NetFlow section with physical and virtual views."),
    ]
    for tier in tiers:
        row = t.add_row().cells
        for i, val in enumerate(tier):
            row[i].text = val

    add_heading(doc, "4.4 Technical Approach", level=2)
    add_para(
        doc,
        "NetFlow collection is deployed as a distributed system with two placements, chosen per site "
        "based on operational preference and network topology:",
    )
    add_bullet(
        doc,
        "Central collector — a hosted BackgroundService inside MicroDataCenter-WebAPI that binds a UDP "
        "listener (default 2055 for v5/v9, 4739 for IPFIX), parses records, and writes them to the "
        "central store. Suitable when exporters can reach the central backend directly.",
    )
    add_bullet(
        doc,
        "Edge collector — the same collector binary, packaged to run on an adjacent MDC device at a "
        "site. It keeps its own local storage and exposes a small read-only query API back to the "
        "central backend. Suitable when flow volume should not cross site boundaries, or when "
        "exporters are not routable to the central backend.",
    )
    add_para(
        doc,
        "The central backend acts as a federation layer: when a user or administrator queries flows, "
        "a FlowQueryCoordinator fans the query out to the relevant collectors (by workspace → site "
        "mapping), merges the results, applies tier-based filtering, and returns a single response. "
        "The user-facing API is unchanged whether a given record lives in central or edge storage.",
    )
    add_bullet(doc, "Records are tagged at ingest with exporter ID, observation-point (interface index), and observed src/dst IPs.")
    add_bullet(doc, "A resolver enriches records with VirtualMachineEntry.Id, WorkspaceEntry.Id, and a physical-vs-virtual flag using the existing VM inventory.")
    add_bullet(doc, "A \"top talkers\" WebSocket stream is published both at the workspace scope (user dashboard) and the global scope (admin dashboard).")

    add_heading(doc, "4.5 New Backend Components", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Responsibility"
    for c in hdr:
        set_cell_bg(c, "1F4E79")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    add_kv_row(t, "NetFlowListenerService", "Hosted BackgroundService owning the UDP socket and dispatch loop. Deployable either in-process (central) or as a standalone service on an edge device.")
    add_kv_row(t, "NetFlowV5Parser / V9Parser / IPFIXParser", "Protocol-specific decoders; IPFIX/v9 share the template cache abstraction.")
    add_kv_row(t, "FlowEnrichmentService", "Resolves src/dst IP to VirtualMachineEntry.Id, WorkspaceEntry.Id, virtual-network-id, and physical-vs-virtual interface flag.")
    add_kv_row(t, "FlowRecordSink", "Batched writer to local storage; backpressure-aware. Implementation varies between central and edge deployments.")
    add_kv_row(t, "FlowAggregationService", "Produces rolled-up views (top-N by src/dst/port, per-VM, per-interface, per-workspace, per-minute/hour).")
    add_kv_row(t, "EdgeCollectorQueryApi", "Lightweight authenticated API on each edge collector that the central backend calls to run scoped queries locally and return results.")
    add_kv_row(t, "FlowQueryCoordinator", "Central-side federation layer. Fans queries across the relevant collectors (central + edges), merges results, and applies tier-based filtering before returning.")
    add_kv_row(t, "FlowsController (OData)", "User-tier and admin-tier read endpoints. The backend injects the workspace scope into every non-admin query; admins get an unscoped variant with physical-interface records included.")
    add_kv_row(t, "FlowExporterController", "CRUD for the inventory of exporters, including their home collector (central or a named edge) and IP allow-list.")

    add_heading(doc, "4.6 Storage Model", level=2)
    add_para(
        doc,
        "Given that expected NetFlow volume on the central backend is very small, the existing "
        "Entity Framework Core / relational store is sufficient for central-tier records. No new "
        "time-series database is introduced on the central system. Retention and rollup policies "
        "are configurable so that historical detail can be compacted over time.",
    )
    add_para(
        doc,
        "Adjacent devices that run their own collector keep their own local storage. The local store "
        "can be a file-based embedded database (SQLite) or a small relational instance, depending on "
        "the device profile. Local storage is bounded by a configurable retention window; older "
        "records are compacted to rollups before expiry, and any rollups the central backend "
        "subscribes to are pushed upstream on a schedule.",
        italic=False,
    )
    add_bullet(doc, "Central: EF Core in the existing database. No new dependency.")
    add_bullet(doc, "Edge: per-device local store, chosen to match the device's footprint (SQLite by default).")
    add_bullet(doc, "Federation: queries that span both are handled transparently by FlowQueryCoordinator.")
    add_bullet(doc, "Retention: configurable per-tier (e.g. 24 hours of raw records, 30 days of 1-minute rollups, 12 months of 1-hour rollups).")

    add_heading(doc, "4.7 Frontend Changes", level=2)
    add_para(doc, "Two distinct surfaces, reflecting the two access tiers:", bold=True)

    add_para(doc, "4.7.1 MDC-web-dashboard (Workspace User)", bold=True)
    add_bullet(doc, "New Network tab inside each workspace view.")
    add_bullet(doc, "Per-interface list — every virtual NIC belonging to the workspace's VMs, with live throughput and a sparkline.")
    add_bullet(doc, "Per-interface drilldown — top talkers, protocol and port breakdown, time-range selector, CSV export of the current view.")
    add_bullet(doc, "No references, links, or counts that leak information about other workspaces or physical interfaces.")

    add_para(doc, "4.7.2 MDC-admin-frontend (Administrator)", bold=True)
    add_bullet(doc, "New top-level NetFlow section with subviews for Virtual (per workspace / per interface) and Physical (per host / per uplink).")
    add_bullet(doc, "Exporter inventory — list of devices sending flows, last-seen timestamps, record counts, home collector (central vs named edge).")
    add_bullet(doc, "Global flow explorer — unscoped filterable table with source, destination, protocol, port, volume, and time-range.")
    add_bullet(doc, "Physical interface dashboard — throughput and top talkers on each host NIC (uplink, management, trunks).")
    add_bullet(doc, "Per-VM drilldown — joins flow records to VirtualMachineEntry by IP to answer \"what is this VM talking to?\".")

    add_heading(doc, "4.8 Security & Workspace Isolation", level=2)
    add_bullet(doc, "Every NetFlow read endpoint derives a tier from the caller's Entra ID claims: Administrator tier or Workspace-member tier. There is no unscoped user-tier endpoint.")
    add_bullet(doc, "User-tier queries are rewritten server-side to include a WorkspaceId IN (user's memberships) filter and to exclude any record flagged as physical-interface.")
    add_bullet(doc, "Admin-tier endpoints are gated behind the existing Administrator role claim; enforcement happens in the backend, not just in the UI.")
    add_bullet(doc, "Exporters must be registered and IP-allow-listed at their home collector (central or edge); unsolicited UDP packets are dropped.")
    add_bullet(doc, "Edge collector query APIs authenticate the central backend by mutual TLS or a shared secret rotated via the existing endpoint-registration channel.")
    add_bullet(doc, "Flow records are treated as sensitive metadata and are subject to existing audit logging.")

    # ---- 5. Part B: MeshCentral ----
    add_heading(doc, "5. Part B — MeshCentral Remote Connectivity", level=1)

    add_heading(doc, "5.1 What MeshCentral Provides", level=2)
    add_para(
        doc,
        "MeshCentral is an open-source remote management platform (documentation: "
        "https://docs.meshcentral.com/). It consists of a self-hosted server and a lightweight agent "
        "installed on each managed machine. Once an agent connects, the server offers browser-based "
        "access to that machine over WebSockets — including full desktop (RDP/VNC), terminal (SSH / "
        "PowerShell / shell), file transfer, and power control. Crucially, it works without exposing "
        "the target machine to the internet: the agent dials out to the MeshCentral server, and the "
        "browser connects to the server, which relays the session.",
    )

    add_heading(doc, "5.2 Business Value", level=2)
    add_bullet(doc, "End-users reach their VMs from any browser, on any OS, without VPN client installation.")
    add_bullet(doc, "Administrators gain a single pane of glass for break-glass access to Windows and Linux VMs.")
    add_bullet(doc, "Reduced support overhead — no per-user VPN provisioning, no firewall exceptions for inbound RDP/SSH.")
    add_bullet(doc, "Session recording and audit trail capability (native MeshCentral feature) satisfies compliance use-cases.")

    add_heading(doc, "5.3 Technical Approach", level=2)
    add_para(
        doc,
        "MeshCentral will be deployed as a self-hosted server inside the client's infrastructure. No "
        "public or third-party MeshCentral service is used. MeshCentral is additive: the existing "
        "Proxmox VNC WebSocket relay remains available for every VM. On each VM detail page, users "
        "will see two options — \"Remote Connect (MeshCentral)\" for full desktop, terminal, and "
        "file-transfer sessions, and \"Console (VNC)\" for the low-level hypervisor console that does "
        "not require an in-guest agent.",
    )
    add_para(
        doc,
        "The MDC backend acts as a trust broker: it authenticates the user via Entra ID (the existing "
        "flow), checks that the user is authorised for the target VM, and then mints a short-lived "
        "MeshCentral login token on the user's behalf. The browser is redirected to the self-hosted "
        "MeshCentral web console with that token. The user never sees, and the browser never receives, "
        "MeshCentral's administrative credentials.",
    )

    add_heading(doc, "5.4 New Backend Components", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Responsibility"
    for c in hdr:
        set_cell_bg(c, "1F4E79")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    add_kv_row(t, "MeshCentralClient", "Typed C# wrapper over the MeshCentral control WebSocket (/control.ashx) and REST surface.")
    add_kv_row(t, "MeshCentralService", "Business operations: device lookup, group management, session-token minting, agent health polling.")
    add_kv_row(t, "MeshDeviceLink (EF entity)", "Mapping table linking VirtualMachineEntry.Id to the MeshCentral NodeId; maintained on enrollment.")
    add_kv_row(t, "MeshController (REST)", "Endpoints: POST /session for token mint, GET /status for agent health, POST /enroll for manual link.")
    add_kv_row(t, "MeshEnrollmentHostedService", "Watches VM lifecycle events and ensures every live VM has a corresponding MeshDeviceLink.")

    add_heading(doc, "5.5 Agent Deployment Strategy", level=2)
    add_para(
        doc,
        "For the MeshCentral agent to reach a VM, it must be installed. Two options, chosen with the client:",
    )
    add_bullet(doc, "Bake the agent into the Proxmox VM template. All new VMs created from that template come pre-enrolled. Simplest, but requires re-templating and a plan for existing VMs.")
    add_bullet(doc, "Cloud-init first-boot install. The MDC backend injects a one-time installer command via cloud-init user-data at VM creation. Works for new VMs without template changes; requires a network path from the VM to the MeshCentral server at first boot.")
    add_para(
        doc,
        "A hybrid is likely: cloud-init for new VMs during rollout, template bake once the approach is stable, "
        "and a one-shot remote installer (delivered via the existing VNC console if needed) for legacy VMs.",
        italic=True,
    )

    add_heading(doc, "5.6 Authentication Flow (Summary)", level=2)
    add_bullet(doc, "1. User signs into MDC dashboard with Entra ID (unchanged).")
    add_bullet(doc, "2. User clicks \"Remote Connect\" on a VM detail page.")
    add_bullet(doc, "3. Frontend calls POST /api/mesh/vms/{id}/session with the existing Bearer token.")
    add_bullet(doc, "4. Backend verifies the user's authorisation for that VM (existing permission model).")
    add_bullet(doc, "5. Backend asks MeshCentral — using a server-side service credential — to issue a short-lived login URL scoped to that single device.")
    add_bullet(doc, "6. Backend returns the signed URL to the frontend, which opens it in a new tab or embedded iframe.")
    add_bullet(doc, "7. MeshCentral validates the token, opens the browser-to-agent session, and the user is in.")

    add_heading(doc, "5.7 Frontend Changes", level=2)
    add_bullet(doc, "MDC-web-dashboard — on every VM detail page, surface both options side by side: \"Remote Connect (MeshCentral)\" and the existing \"Console (VNC)\". Neither replaces the other; users choose based on need.")
    add_bullet(doc, "MDC-admin-frontend — add a MeshCentral management section: agent health overview, bulk re-enrollment, device groups, session audit log browser, and a view indicating which VMs have a live agent versus VNC-only reachability.")

    add_heading(doc, "5.8 Security Considerations", level=2)
    add_bullet(doc, "Service credentials for MeshCentral live only on the backend; never transmitted to the browser.")
    add_bullet(doc, "Session tokens are short-lived (recommended: 60 seconds to redeem, session bound thereafter).")
    add_bullet(doc, "MeshCentral supports TOTP 2FA for administrator roles — recommended for admin access paths.")
    add_bullet(doc, "All session activity is logged by MeshCentral and surfaced in the admin audit view.")
    add_bullet(doc, "The MeshCentral server is placed behind TLS with a certificate from the client's PKI; agent-server traffic is TLS-pinned.")

    # ---- 6. Part C: Proxmox Observability via PDM ----
    add_heading(doc, "6. Part C — Proxmox Observability via Proxmox Datacenter Manager", level=1)

    add_heading(doc, "6.1 What Proxmox Datacenter Manager Provides", level=2)
    add_para(
        doc,
        "Proxmox Datacenter Manager (PDM) is Proxmox Server Solutions' central management tool for "
        "operating multiple Proxmox VE clusters as a single estate "
        "(https://www.proxmox.com/en/products/proxmox-datacenter-manager/overview). A single PDM "
        "instance discovers its registered Proxmox clusters and nodes, collects and retains health "
        "and utilisation telemetry for each, and exposes both a web UI and a REST API spanning the "
        "entire fleet. It handles the edge-to-centre aggregation that the MDC platform would "
        "otherwise have to build from scratch.",
    )
    add_bullet(doc, "Fleet inventory — clusters, nodes, VMs, containers, storage pools, network bridges.")
    add_bullet(doc, "Telemetry — CPU, memory, storage I/O and utilisation, network throughput, uptime, at node, VM, and container granularity.")
    add_bullet(doc, "Event and task history — who did what on which cluster, when.")
    add_bullet(doc, "A REST API and an authentication model (API tokens) suitable for machine consumption.")

    add_heading(doc, "6.2 Why PDM Simplifies This Work", level=2)
    add_para(
        doc,
        "Building multi-cluster telemetry aggregation in the MDC backend would duplicate significant "
        "infrastructure that Proxmox already ships. Delegating that concern to PDM means the MDC "
        "backend treats telemetry as a read-through cache over a single HTTPS endpoint, regardless "
        "of how many Proxmox clusters or sites are in play. Adding a new cluster becomes an "
        "operational step in PDM rather than a code change in MDC.",
    )
    add_bullet(doc, "Distributed collection is solved upstream — MDC does not operate per-site agents for telemetry.")
    add_bullet(doc, "Telemetry retention and rollups are PDM's responsibility.")
    add_bullet(doc, "MDC remains the authoritative tenancy layer (workspaces, users, permissions); PDM remains the authoritative infrastructure layer. Each stays in its lane.")

    add_heading(doc, "6.3 Business Value", level=2)
    add_bullet(doc, "Workspace users see live health of their own VMs without waiting for a support escalation.")
    add_bullet(doc, "Administrators diagnose cluster- and host-level issues from the same portal they use for everything else — no context switch to the PDM UI for day-to-day work.")
    add_bullet(doc, "Single operational dashboard correlating tenancy (MDC), traffic (NetFlow), remote access (MeshCentral + VNC), and health (PDM) in one place.")
    add_bullet(doc, "Lower operational burden adding new clusters or sites — registering them in PDM is enough for them to appear in MDC.")

    add_heading(doc, "6.4 Access Tiers", level=2)
    add_para(
        doc,
        "Observability visibility mirrors the NetFlow tiering model so that users see a consistent "
        "boundary across features.",
    )
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Audience", "Scope of Visibility", "Where Surfaced"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tiers = [
        ("Workspace User",
         "Health and utilisation metrics for VMs and containers belonging to workspaces the user is a member of. No visibility into host-level metrics, other workspaces, or cluster internals.",
         "MDC-web-dashboard — Observability panel inside each VM detail page and a workspace-scoped summary view."),
        ("Administrator",
         "Full fleet view: every cluster, node, VM, container, storage pool. Host-level CPU/memory/IO/network, cluster quorum state, storage health, PDM events and tasks.",
         "MDC-admin-frontend — new Observability section with fleet, cluster, node, and storage subviews."),
    ]
    for tier in tiers:
        row = t.add_row().cells
        for i, val in enumerate(tier):
            row[i].text = val

    add_heading(doc, "6.5 Technical Approach", level=2)
    add_para(
        doc,
        "A new provider module is added to the .NET backend under "
        "MDC.Core/Services/Providers/ProxmoxDatacenterManager/. It wraps the PDM REST API with a "
        "typed client, translates PDM resource identifiers into MDC identifiers (workspace, site, "
        "VM), and serves a read-through cache to the frontends. The MDC backend does not store PDM "
        "data long-term; PDM is the source of truth.",
    )
    add_bullet(doc, "Authentication: a dedicated PDM API token held only on the MDC backend; never exposed to the browser.")
    add_bullet(doc, "Resource mapping: a lightweight table linking MDC VirtualMachineEntry.Id to PDM cluster/node/vmid, populated at VM creation and reconciled on a schedule.")
    add_bullet(doc, "Read-through cache: short-TTL in-memory cache (configurable, default 5 seconds for live metrics, 60 seconds for inventory) to keep PDM load bounded under heavy portal use.")
    add_bullet(doc, "Live updates: optional WebSocket push to connected dashboards using the existing WebSocket infrastructure; backed by the cache, not by direct PDM polling per client.")
    add_bullet(doc, "Tier filtering: all user-tier queries are rewritten server-side to constrain the PDM response to the caller's workspace VMs before returning.")

    add_heading(doc, "6.6 New Backend Components", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Responsibility"
    for c in hdr:
        set_cell_bg(c, "1F4E79")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    add_kv_row(t, "PdmClient", "Typed C# wrapper over the PDM REST API (authentication, inventory, metrics, events).")
    add_kv_row(t, "PdmResourceResolver", "Translates between MDC identifiers (WorkspaceId, VirtualMachineEntry.Id) and PDM identifiers (cluster/node/vmid).")
    add_kv_row(t, "PdmTelemetryCache", "Short-TTL in-memory cache shielding PDM from per-client query storms; single-flight to collapse concurrent requests.")
    add_kv_row(t, "PdmReconciliationService", "Background job that refreshes the resource-mapping table and surfaces drift (PDM VMs not known to MDC, or vice versa).")
    add_kv_row(t, "ObservabilityController", "User-tier and admin-tier read endpoints for metrics, inventory, and events. User-tier queries are scoped to workspace membership server-side.")
    add_kv_row(t, "PdmHealthCheck", "Reports PDM reachability to the existing health endpoints and drives the degraded-mode banner in the UI.")

    add_heading(doc, "6.7 Data Freshness, Caching, and Degraded Mode", level=2)
    add_bullet(doc, "Live-metric endpoints return data no older than the cache TTL (default 5 seconds); the UI shows the freshness timestamp.")
    add_bullet(doc, "Inventory endpoints use a longer TTL (default 60 seconds) and are invalidated by MDC-initiated VM lifecycle events.")
    add_bullet(doc, "If PDM is unreachable, observability views enter a clearly-indicated degraded mode: last-known cached values are shown, annotated with staleness, rather than a blank page.")
    add_bullet(doc, "Loss of PDM does not affect any other MDC functionality — tenancy, VM management via Proxmox directly, NetFlow, MeshCentral, and VNC all continue.")

    add_heading(doc, "6.8 Frontend Changes", level=2)
    add_para(doc, "Two distinct surfaces, matching the tiering model:", bold=True)

    add_para(doc, "6.8.1 MDC-web-dashboard (Workspace User)", bold=True)
    add_bullet(doc, "New Observability panel on each VM detail page — CPU, memory, disk I/O, network throughput, uptime, with sparkline and short history.")
    add_bullet(doc, "Workspace summary page — at-a-glance health of every VM in the workspace (green/amber/red based on configurable thresholds).")
    add_bullet(doc, "Integration with existing VM detail page so network flows (NetFlow), remote access (MeshCentral/VNC), and health (PDM) appear in a single coherent view.")

    add_para(doc, "6.8.2 MDC-admin-frontend (Administrator)", bold=True)
    add_bullet(doc, "New top-level Observability section with subviews for Fleet, Clusters, Nodes, Storage, and Events.")
    add_bullet(doc, "Fleet view — every cluster, quorum state, node count, aggregate resource utilisation.")
    add_bullet(doc, "Node view — per-host CPU/memory/IO/network, hosted VM list with drilldown.")
    add_bullet(doc, "Storage view — pool capacity, usage trend, and health per cluster.")
    add_bullet(doc, "Events view — PDM task and event log, filterable by cluster, node, or VM.")
    add_bullet(doc, "Clear indicator of PDM reachability state at the top of the section.")

    add_heading(doc, "6.9 Security & Workspace Isolation", level=2)
    add_bullet(doc, "The PDM API token lives only on the MDC backend (in the secrets store used for existing provider credentials). It is never delivered to the browser, embedded in client-side code, or logged.")
    add_bullet(doc, "The PDM instance is not exposed to end-users directly; the only path from browser to PDM is via MDC backend endpoints, which enforce MDC's own authentication and authorisation.")
    add_bullet(doc, "User-tier observability queries are rewritten server-side to include the caller's workspace-VM list before the PDM request is issued; the PDM response is filtered a second time defensively before returning to the client.")
    add_bullet(doc, "Admin-tier endpoints are gated behind the existing Administrator role claim.")
    add_bullet(doc, "PDM access is read-mostly; write operations (VM lifecycle) continue to use the existing Proxmox control paths, keeping the blast radius of any PDM credential compromise bounded.")

    # ---- 7. Phased Delivery Plan ----
    add_heading(doc, "7. Phased Delivery Plan", level=1)
    add_para(
        doc,
        "Delivery is organised as six sequential phases. Each phase is independently reviewable and "
        "produces a demonstrable artefact. Durations are engineering estimates assuming one full-time "
        "engineer and will be refined after the open questions in section 9 are answered.",
    )
    add_phase_table(doc, [
        ("Phase 1", "PDM stand-up (self-hosted) and backend PDM provider thin slice — auth, inventory read, resource mapping for one cluster.",
         "1.5 weeks", "Backend can list Proxmox clusters/nodes/VMs via PDM and resolve them to MDC workspaces."),
        ("Phase 2", "Workspace-user Observability panel in MDC-web-dashboard — per-VM live metrics scoped to the caller's workspaces.",
         "1.5 weeks", "Workspace user sees live CPU/memory/network/IO for their own VMs, no leakage across workspaces."),
        ("Phase 3", "Admin Observability section in MDC-admin-frontend — fleet, cluster, node, storage, events subviews with degraded-mode indicator.",
         "2 weeks", "Administrator can operate the Proxmox estate from the MDC portal without switching to PDM UI for routine work."),
        ("Phase 4", "Self-hosted MeshCentral server stand-up and backend trust-broker provider thin slice (one VM, end-to-end).",
         "1.5 weeks", "A single enrolled VM reachable via \"Remote Connect\" from the user dashboard, shown alongside the existing VNC option."),
        ("Phase 5", "User dashboard Remote Connect UI (coexisting with VNC) and agent enrollment automation for new VMs.",
         "1.5 weeks", "Any newly created VM auto-enrolls and is reachable from the browser; VNC option remains untouched."),
        ("Phase 6", "NetFlow central collector, enrichment, workspace / VM resolution, and EF Core persistence with tier-aware query API.",
         "2 weeks", "Flows from at least one exporter visible via backend API, correctly tagged with workspace and interface."),
        ("Phase 7", "Workspace-user NetFlow UI in MDC-web-dashboard — per-interface views scoped to the caller's workspaces.",
         "1.5 weeks", "Workspace user can see per-vNIC flow analysis for their VMs, with no leakage across workspaces."),
        ("Phase 8", "Admin NetFlow UI in MDC-admin-frontend — global flow explorer, physical-interface dashboard, exporter inventory, per-VM drilldown.",
         "2 weeks", "Administrator can answer \"what is VM X talking to\" and monitor physical uplinks from a single place."),
        ("Phase 9", "Edge NetFlow collector package and federation — deployable to adjacent site devices, with FlowQueryCoordinator on the central side.",
         "1.5 weeks", "At least one edge collector in production at a client site, queryable transparently from the central UIs."),
        ("Phase 10", "Admin MeshCentral management UI — health, groups, session audit.",
         "1.5 weeks", "Full admin lifecycle operations for MeshCentral available without leaving MDC admin."),
        ("Phase 11", "Hardening — failure-mode rehearsal, load validation, documentation, handover, client training.",
         "1 week", "Production-ready release with runbooks and a recorded training session."),
    ])
    add_para(
        doc,
        "Sequencing note: PDM-based observability is delivered first because it is the lowest-risk "
        "integration (PDM handles the hard parts upstream) and immediately makes the existing VM "
        "inventory more useful across both portals.",
        italic=True,
    )

    # ---- 8. Prerequisites From Client ----
    add_heading(doc, "8. Prerequisites From the Client", level=1)
    add_para(doc, "To keep the schedule on track, the following are required before or during early phases:")
    add_bullet(doc, "A deployed PDM instance with every Proxmox cluster to be observed already registered in it; a hostname and TLS certificate for the PDM web/API endpoint.")
    add_bullet(doc, "A dedicated PDM API token with read-only scope across the relevant clusters, issued to the MDC backend service account.")
    add_bullet(doc, "Confirmation of the PDM telemetry retention window so MDC UI time-range selectors can be bounded accordingly.")
    add_bullet(doc, "A hostname and TLS certificate for the self-hosted MeshCentral server (e.g. mesh.client-domain.tld).")
    add_bullet(doc, "A VM or container host on which to deploy the MeshCentral server (modest spec: 2 vCPU, 4 GB RAM, 50 GB disk for a small fleet).")
    add_bullet(doc, "Network rules allowing outbound HTTPS/WSS from managed VMs to the MeshCentral server.")
    add_bullet(doc, "Designated NetFlow exporters for virtual traffic (the virtual switch / bridge / ZeroTier exporter on each host) and for physical traffic (top-of-rack switches, firewalls, uplink routers). Permission to enable flow export on each.")
    add_bullet(doc, "For each site that should run a local edge collector: identification of the adjacent device that will host it, and confirmation of its network reach to the local exporters.")
    add_bullet(doc, "Mapping of exporter observation-points (interface indices) to their physical or virtual identity, so the UI can label them meaningfully.")
    add_bullet(doc, "Approval for the Administrator role claim name and membership process (if different from what already exists in Entra ID), and confirmation of the workspace-membership claim used for user-tier scoping.")
    add_bullet(doc, "A test VM and a test workspace on which all three integrations can be validated before fleet rollout.")

    # ---- 9. Risks & Mitigations ----
    add_heading(doc, "9. Risks & Mitigations", level=1)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Risk", "Impact", "Mitigation"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    risks = [
        ("NetFlow volume exceeds storage capacity", "Data loss, query slowness",
         "Abstract the storage layer; confirm peak rate before go-live; enable downsampling and retention policies."),
        ("MeshCentral server becomes a single point of failure for remote access",
         "Users unable to reach VMs during outage",
         "Retain existing VNC console as fallback; document manual recovery runbook; optionally deploy MeshCentral in HA pair."),
        ("Agent installation fails on legacy VMs", "Incomplete coverage",
         "Provide a manual installer path and a clearly documented \"unmanaged VM\" state in the UI."),
        ("NetFlow exporters send malformed or spoofed packets",
         "Parser crash, false data", "IP allow-list, strict parser with bounded buffers, fuzz-testing of parsers before release."),
        ("Entra ID token to MeshCentral trust-broker gap misused",
         "Unauthorised access to a VM", "Server-side authorisation check before every session mint; short-lived tokens; audit every mint."),
        ("Clock skew between collector and exporters",
         "Misleading time-series views", "Enforce NTP on collector host; display timestamps in exporter-local and UTC side by side."),
        ("PDM outage or API unavailability",
         "Observability views become stale or empty",
         "Read-through cache serves last-known values during short outages; UI displays a clearly-indicated degraded-mode banner with staleness timestamps; no dependency on PDM for MDC's core VM management paths."),
        ("PDM product maturity (relatively new Proxmox offering)",
         "API surface may change between PDM versions",
         "Isolate all PDM interaction behind PdmClient and a stable internal interface; pin to a tested PDM version; run compatibility check in CI against the target PDM version."),
        ("PDM API token compromise",
         "Read access to fleet telemetry by a third party",
         "Token scoped read-only; stored in the existing secrets store; rotated on a schedule; write operations to Proxmox continue to use separate control paths so a telemetry token cannot be escalated into fleet changes."),
    ]
    for risk in risks:
        row = t.add_row().cells
        for i, val in enumerate(risk):
            row[i].text = val

    # ---- 10. Open Questions ----
    add_heading(doc, "10. Open Questions for the Client", level=1)
    add_para(doc, "Answers to these questions will let us lock scope and sharpen estimates:")
    add_bullet(doc, "Is PDM already deployed, or is deployment part of this engagement? Which version, and on what host?")
    add_bullet(doc, "Which Proxmox clusters must be registered in PDM for Phase 1, and which are planned to follow?")
    add_bullet(doc, "PDM telemetry retention window — default, or any policy override?")
    add_bullet(doc, "Should the MDC portal link out to the PDM UI for advanced admin tasks, or must every interaction stay inside the MDC portal?")
    add_bullet(doc, "For each site: should NetFlow exporters send to the central collector, to a local edge collector, or a mix? Which adjacent device at each site will host the edge collector?")
    add_bullet(doc, "Is there an existing flow collector anywhere on the network whose output we should consume rather than duplicate?")
    add_bullet(doc, "Retention policy for flow records: how long must raw records be kept, and how long must rolled-up aggregates be kept? Same policy for edge and central, or different?")
    add_bullet(doc, "At the user tier, should flow and observability visibility be available to every workspace member, or only to specified workspace roles (owner / admin / operator)?")
    add_bullet(doc, "Are there any workspaces or interfaces that should be excluded from user-tier visibility entirely (e.g. sensitive internal services)?")
    add_bullet(doc, "Is session recording of MeshCentral sessions a hard compliance requirement? For admin sessions only, or for all sessions?")
    add_bullet(doc, "Preferred deployment topology for MeshCentral: single node, HA pair, or regional nodes close to user populations?")
    add_bullet(doc, "Labelling convention for physical interfaces in the admin UI (by host-and-ifindex, by rack-location, by role, or some combination)?")

    # ---- 11. Deliverables ----
    add_heading(doc, "11. Deliverables", level=1)
    add_bullet(doc, "Source code changes across the three repositories, delivered via reviewable pull requests.")
    add_bullet(doc, "Database migration scripts for the new EF entities.")
    add_bullet(doc, "Deployment artefacts: PDM provider configuration, MeshCentral server configuration, NetFlow collector service files (central and edge packages), updated application deployment manifests.")
    add_bullet(doc, "Operator runbooks: PDM connectivity and credential rotation, NetFlow storage and retention management, MeshCentral backup and recovery, agent re-enrollment.")
    add_bullet(doc, "End-user quick-start guide for \"Remote Connect\" and the observability / flow views.")
    add_bullet(doc, "Administrator guide covering the new admin UI surfaces across all three capabilities.")
    add_bullet(doc, "Test evidence: integration tests, end-to-end MeshCentral session test, PDM-down degraded-mode test, NetFlow parser conformance tests.")

    # ---- 12. Assumptions ----
    add_heading(doc, "12. Assumptions", level=1)
    add_bullet(doc, "The existing Entra ID tenant and MSAL configuration are production-ready and will be reused.")
    add_bullet(doc, "The existing CI/CD pipelines can accommodate new .NET projects and new frontend routes without restructuring.")
    add_bullet(doc, "Proxmox remains the underlying hypervisor for the duration of this engagement, and PDM is the chosen central management surface for the Proxmox estate.")
    add_bullet(doc, "The PDM REST API is used as a read source for telemetry; no writes from MDC to PDM are in scope.")
    add_bullet(doc, "Licensing is not a blocker: MeshCentral is AGPL-licensed and self-hosted; NetFlow is an open protocol with no per-device licensing; PDM licensing follows the client's Proxmox subscription model.")
    add_bullet(doc, "The client accepts that any external hosted services (e.g. MeshCentral.com public instance, Proxmox-hosted PDM) will not be used; all components are self-hosted.")

    # ---- 13. Approval ----
    add_heading(doc, "13. Approval", level=1)
    add_para(doc, "By approving this document, the client confirms the scope and phased plan described above, "
                  "and authorises the engineering team to proceed with Phase 1.")
    doc.add_paragraph()
    t = doc.add_table(rows=3, cols=2)
    t.style = "Light Grid Accent 1"
    labels = [("Client Representative", ""), ("Signature & Date", ""), ("Project Lead (MDC team)", "")]
    for i, (k, v) in enumerate(labels):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v

    out = "/Users/sudipto/Desktop/projects/MDC-web-dashboard/MDC_Integration_Proposal_NetFlow_MeshCentral.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
